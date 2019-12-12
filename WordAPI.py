
########### Generate a offenses report from Qradar ########
########### Author: Nadav Yarkon #############
    ### For Educational Purposes Only ###

#Date: 01.12.18
#Author: Nadav Yarkon
#Email: Nadavy2469 @ gmail.com
#https: https://github.com/Nadav-Yarkon


from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

class DocxAPI:
    __document = None
    __currentPar = None

    def __init__(self):
        self.__document = Document()
        __currentPar = None

    def writeText(self,text,color,size,newLine = False):
        if newLine:
            __currentPar = self.newParagraph()
        fonter = self.__currentPar.add_run(text)
        fonter.font.color.rgb = self.__getColor(color)
        fonter.font.size = Pt(size)

    def newParagraph(self):
        return self.__document.add_paragraph()

    def header(self,text):
        self.__document.add_heading(text, 1)
        self.__currentPar = self.newParagraph()

    def __getColor(self,color):
        if color == 'blue':
            return RGBColor(0x42, 0x50, 0xE9)
        elif color == 'red':
            return RGBColor(0xFF, 0x00, 0x00)
        elif color == 'black':
            return RGBColor(0,0,0)
        elif color == 'none':
            return None

    def saveFile(self,fileName):
        self.__document.save(fileName)
